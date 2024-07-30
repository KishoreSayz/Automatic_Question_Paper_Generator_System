import os
import random
import sqlite3
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog

class QuestionPaperGenerator(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Automatic Question Paper Generator")
        self.geometry("800x600")
        self.create_main_window()

    def create_main_window(self):
        main_frame = ttk.Frame(self, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        
        heading = ttk.Label(main_frame, text="Welcome to Automatic Question Paper Generator for Avanthi Institute of Engineering and Technology",
                            font=("Helvetica", 20, "bold"), foreground="red")
        heading.grid(row=0, column=0, columnspan=2, pady=(50, 20))
        
        creator_info = ttk.Label(main_frame, text="Created by:\n\n 1. K.V.S.S.Kishore (Roll No: 21Q71A0575)\n\n 2. S.V.S.K.Aditya (Roll No: 22Q75A0516)\n\n 3. K.Arjun Sai (Roll No: 21Q71A0568) \n\n Of 4th year CSE",
                                 font=("Helvetica", 12), wraplength=750, justify="left")
        creator_info.grid(row=1, column=0, sticky=tk.W, padx=(20, 0))
        
        next_button = ttk.Button(main_frame, text="Next", command=self.open_subject_window, style='Red.TButton')
        next_button.grid(row=2, column=0, columnspan=2, pady=(20, 0))

        style = ttk.Style()
        style.configure('Red.TButton', foreground='black', background='red', font=('Helvetica', 12, 'bold'))

    def open_subject_window(self):
        subject_window = tk.Toplevel(self)
        subject_window.title("Subject Options")
        subject_window.geometry("600x400")
        
        heading = ttk.Label(subject_window, text="Choose Semester and Subject", font=("Helvetica", 16))
        heading.pack(pady=(20, 30))
        
        ttk.Label(subject_window, text="Semester:").pack(anchor=tk.W, padx=20)
        self.semester = ttk.Combobox(subject_window, values=["1-1", "1-2", "2-1", "2-2", "3-1", "3-2", "4-1"])
        self.semester.pack(anchor=tk.W, padx=20, pady=(0, 20))
        self.semester.bind("<<ComboboxSelected>>", self.update_subjects)

        ttk.Label(subject_window, text="Subject:").pack(anchor=tk.W, padx=20)
        self.subject = ttk.Combobox(subject_window)
        self.subject.pack(anchor=tk.W, padx=20, pady=(0, 20))
        
        generate_auto_button = ttk.Button(subject_window, text="Generate Automatically", command=self.generate_automatically)
        generate_auto_button.pack(pady=(20, 10))

        manual_entry_button = ttk.Button(subject_window, text="Generate Manually", command=self.manual_question_entry)
        manual_entry_button.pack(pady=(0, 20))

        change_subject_button = ttk.Button(subject_window, text="Change Subject", command=subject_window.destroy)
        change_subject_button.pack()

    def update_subjects(self, event):
        semester_subjects = {
            "1-1": ["M1", "ENG", "Physics", "Python"],
            "1-2": ["M2", "PPS", "Chemistry", "DS", "CO"],
            "2-1": ["M3", "MFCS", "C++", "OS", "SE"],
            "2-2": ["P&S", "DBMS", "Java", "FLAT", "MEFA"],
            "3-1": ["DWDM", "CN", "SPM", "DAA", "DLD"],
            "3-2": ["ML", "CD", "CNS", "OOAD", "SC"],
            "4-1": ["IAML", "IoT", "CS&F", "DLT", "WNS", "UHV"]
        }
        selected_semester = self.semester.get()
        self.subject['values'] = semester_subjects.get(selected_semester, [])
        self.subject.set("")

    def manual_question_entry(self):
        subject = self.subject.get()
        questions = []
        question_number = 1

        while len(questions) < 9:
            question = simpledialog.askstring("Input", f"Enter question {question_number} for {subject}:")
            if question:
                questions.append(question)
                question_number += 1
            else:
                messagebox.showwarning("Warning", "Please enter all 9 questions.")
                return
        
        # Ensure no duplicates
        questions = list(dict.fromkeys(questions))
        
        # If there are still duplicates, fill with dummy questions
        while len(questions) < 9:
            questions.append("Dummy question to fill space.")

        questions_set_1 = questions[:3]
        questions_set_2 = questions[3:6]
        questions_set_3 = questions[6:9]

        self.generate_question_paper(subject, questions_set_1, questions_set_2, questions_set_3)

        messagebox.showinfo("Success", "Thank you for using Automatic Question Paper Generator.\n\n Your question paper is saved in the 'Generated papers' folder.")

    def generate_automatically(self):
        semester = self.semester.get()
        subject = self.subject.get()

        if not semester or not subject:
            messagebox.showwarning("Warning", "Please select both semester and subject.")
            return
        
        num_questions = 9
        questions = self.get_random_questions(subject, num_questions)

        if len(questions) < num_questions:
            messagebox.showerror("Error", f"Not enough questions found for the subject: {subject}")
            return
        
        question_ids = [q[0] for q in questions]
        question_texts = [q[1] for q in questions]

        question_texts = list(dict.fromkeys(question_texts))

        while len(question_texts) < num_questions:
            additional_questions = self.get_random_questions(subject, num_questions - len(question_texts))
            question_ids.extend([q[0] for q in additional_questions])
            question_texts.extend([q[1] for q in additional_questions])
            question_texts = list(dict.fromkeys(question_texts))

        questions_set_1 = question_texts[:3]
        questions_set_2 = question_texts[3:6]
        questions_set_3 = question_texts[6:9]

        self.generate_question_paper(subject, questions_set_1, questions_set_2, questions_set_3, question_ids)

        messagebox.showinfo("Success", "Thank you for using Automatic Question Paper Generator.\nYour question paper is saved in the 'Generated papers' folder in the 'Downloads' folder of your system.")

    def get_random_questions(self, subject, num_questions):
        conn = sqlite3.connect(self.get_db_path())
        cursor = conn.cursor()

        # Get the IDs of the questions already used
        cursor.execute('SELECT question_id FROM used_questions')
        used_question_ids = [row[0] for row in cursor.fetchall()]

        # Select random questions excluding the used ones
        cursor.execute('''
        SELECT id, question FROM questions 
        WHERE subject = ? AND id NOT IN ({seq})
        ORDER BY RANDOM() LIMIT ?
        '''.format(seq=','.join(['?']*len(used_question_ids))), [subject, *used_question_ids, num_questions])

        questions = cursor.fetchall()
        conn.close()

        return questions

    def get_db_path(self):
        base_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(base_dir, 'qp.db')
        print(f"Database path: {db_path}")  # Debugging line
        return db_path

    def generate_question_paper(self, subject, questions_set_1, questions_set_2, questions_set_3, question_ids=None):
        # Load the template Word file
        base_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(base_dir, 'Template', 'Template.docx')
        print(f"Template path: {template_path}")  # Debugging line
        try:
            document = Document(template_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load template: {e}")
            return

        # Insert questions into the document
        def insert_questions(table, questions):
            for i, question in enumerate(questions):
                row = table.rows[i + 1]  # Assuming first row is header
                row.cells[1].text = question  # Insert question text into the second cell

        # Find and update tables for each set of questions
        tables = document.tables
        insert_questions(tables[0], questions_set_1)  # First page
        insert_questions(tables[1], questions_set_2)  # Second page
        insert_questions(tables[2], questions_set_3)  # Third page

        # Insert subject name and code in each table of the document
        subject_codes = {
            'dwdm': '101',
            'cn': '102',
            'spm': '103',
            'daa': '104',
            'dld': '105',
            'iaml': '106'
            # Add other subjects and codes here
        }
        subject_code = subject_codes.get(subject.lower(), 'UNKNOWN')

        for paragraph in document.paragraphs:
            if 'Subject:' in paragraph.text:
                paragraph.text = paragraph.text.replace('Subject:', f'Subject: {subject.upper()}')
            if 'Subject Code:' in paragraph.text:
                paragraph.text = paragraph.text.replace('Subject Code:', f'Subject Code: {subject_code}')

        # Save the document
        output_dir = os.path.expanduser('~/Downloads/Generated papers')
        print(f"Output directory: {output_dir}")  # Debugging line
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        try:
            if question_ids:
                with open(os.path.join(output_dir, f'{subject}_{timestamp}.txt'), 'w') as file:
                    file.write('\n'.join(map(str, question_ids)))
            document.save(os.path.join(output_dir, f'{subject}_{timestamp}.docx'))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document: {e}")
            return

        # Update used_questions table in the database
        if question_ids:
            conn = sqlite3.connect(self.get_db_path())
            cursor = conn.cursor()
            placeholders = ','.join(['(?)']*len(question_ids))
            try:
                cursor.execute(f'INSERT INTO used_questions (question_id) VALUES {placeholders}', question_ids)
                conn.commit()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to update database: {e}")
            finally:
                conn.close()

        self.destroy()

if __name__ == "__main__":
    app = QuestionPaperGenerator()
    app.mainloop()
